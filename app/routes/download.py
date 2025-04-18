from flask import Blueprint, request, send_file
from app.database import db
from app.models import Patient
from docx import Document
import os

download_bp = Blueprint('download', name)

@download_bp.route('/download/recruit/<int:recruit_id>', methods=['GET'])
def download_recruit_info(recruit_id):
    # Получаем данные призывника
    recruit = Patient.query.get(recruit_id)
    if not recruit:
        return {"error": "Призывник не найден"}, 404

    # Создаем документ Word
    doc = Document()
    doc.add_heading('Личная информация призывника', level=1)
    doc.add_paragraph(f"ФИО: {recruit.full_name}")
    doc.add_paragraph(f"Дата рождения: {recruit.birth_date}")
    doc.add_paragraph(f"Адрес: {recruit.address}")
    doc.add_paragraph(f"Телефон: {recruit.phone}")
    doc.add_paragraph(f"Email: {recruit.email}")
    doc.add_paragraph(f"Медицинская карта: {recruit.medical_record}")
    doc.add_paragraph(f"Полис ОМС: {recruit.oms_policy}")
    doc.add_paragraph(f"Группа крови: {recruit.blood_group}")
    doc.add_paragraph(f"Хронические заболевания: {recruit.chronic_diseases}")

    # Путь сохранения файла
    file_path = os.path.join("instance", f"recruit_{recruit_id}.docx")
    doc.save(file_path)

    # Отправляем файл пользователю
    return send_file(file_path, as_attachment=True, download_name=f"recruit_{recruit_id}.docx")